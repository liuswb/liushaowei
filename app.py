from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, make_response, session, Response, send_file
from datetime import datetime, timedelta
from work_manager import WorkManager
from urllib.parse import quote
from functools import wraps
import io
import csv
from docx import Document
from docx.shared import Pt
import os
from routes.project_routes import bp as project_bp
from config.district_config import get_grouped_districts

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # 用于flash消息
manager = WorkManager()

# 注册蓝图
app.register_blueprint(project_bp)

# 登录装饰器
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# 权限检查装饰器
def permission_required(module, action):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not manager.check_permission(session['user_id'], module, action):
                flash('权限不足')
                return redirect(url_for('index'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

@app.template_filter('datetime')
def format_datetime(value):
    if isinstance(value, str):
        value = datetime.fromisoformat(value)
    return value.strftime('%Y-%m-%d %H:%M:%S')

@app.route('/')
@login_required
def index():
    if session['role'] == 'admin':
        projects = manager.get_projects_by_activity()
    else:
        projects = manager.get_user_projects(session['user_id'])
    tasks = manager.get_user_tasks(session['user_id'])
    project_stats = manager.get_project_statistics()
    district_groups = get_grouped_districts()
    
    return render_template('index.html', 
                         active_projects=projects['active_projects'],
                         recent_inactive=projects['recent_inactive'],
                         long_inactive=projects['long_inactive'],
                         tasks=tasks,
                         project_stats=project_stats,
                         district_groups=district_groups)

@app.route('/add_project', methods=['POST'])
def add_project():
    if request.method == 'POST':
        try:
            # 获取表单数据
            project_id = request.form.get('project_id', '').strip()
            client_name = request.form.get('client_name', '').strip()
            stage = request.form.get('stage', '').strip()
            status = request.form.get('status', '').strip()
            area = request.form.get('area', '').strip()
            project_manager = request.form.get('manager', '').strip()
            manager_phone = request.form.get('manager_phone', '').strip()
            notes = ''

            # 验证必填字段
            if not all([project_id, client_name, stage, area, project_manager, manager_phone]):
                missing = []
                if not project_id: missing.append('项目ID')
                if not client_name: missing.append('客户名称')
                if not stage: missing.append('项目环节')
                if not area: missing.append('项目区域')
                if not project_manager: missing.append('项目经理')
                if not manager_phone: missing.append('联系电话')
                flash(f'请填写以下必填字段：{", ".join(missing)}')
                return redirect(url_for('index'))

            # 验证项目ID格式
            try:
                project_id = int(project_id)
                if project_id <= 0:
                    flash('项目ID必须为正整数')
                    return redirect(url_for('index'))
            except ValueError:
                flash('项目ID必须为数字')
                return redirect(url_for('index'))

            # 添加项目
            print(f"正在添加项目：ID={project_id}, 客户={client_name}, 环节={stage}")
            manager.add_project(project_id, client_name, stage, status, notes, area, project_manager, manager_phone)
            flash('项目添加成功！')
        except ValueError as e:
            print(f"添加项目时出现 ValueError：{e}")
            flash(str(e))
        except Exception as e:
            print(f"添加项目时出现错误：{e}")
            flash(f'添加项目时出错：{str(e)}')
    return redirect(url_for('index'))

@app.route('/add_task', methods=['POST'])
@login_required
def add_task():
    try:
        project_id = request.form.get('project_id')
        content = request.form.get('content')
        priority = request.form.get('priority')
        
        if not content:
            flash('请输入任务内容')
            return redirect(url_for('index'))
        
        success = manager.add_task(project_id, content, priority, session['user_id'])
        
        if success:
            flash('任务添加成功！')
        else:
            flash('任务添加失败，请重试')
            
    except Exception as e:
        flash(f'添加任务时出错：{str(e)}')
    
    return redirect(url_for('index'))

@app.route('/complete_task/<int:task_id>', methods=['POST'])
def complete_task(task_id):
    completion_note = request.json.get('completion_note', '')
    success = manager.complete_task(task_id, completion_note)
    return jsonify({'success': success})

@app.route('/cancel_task/<int:task_id>', methods=['POST'])
def cancel_task(task_id):
    cancel_reason = request.json.get('cancel_reason', '')
    success = manager.cancel_task(task_id, cancel_reason)
    return jsonify({'success': success})

@app.route('/update_project_state/<int:project_id>', methods=['POST'])
def update_project_state(project_id):
    data = request.get_json()
    new_state = data.get('state')
    if new_state in ['active', 'recent_inactive', 'long_inactive']:
        manager.update_project_state(project_id, new_state)
        return jsonify({'success': True})
    return jsonify({'success': False, 'error': 'Invalid state'})

@app.route('/project_history')
def project_history():
    project_id = request.args.get('project_id')
    if not project_id:
        flash('请提供项目ID')
        return redirect(url_for('index'))
    
    history = manager.get_project_history(project_id=project_id)
    if not history:
        history = []
    
    # 获取项目信息
    project = manager.get_project(project_id)
    if not project:
        flash('项目不存在')
        return redirect(url_for('index'))
    
    # 获取设备信息
    devices = manager.get_project_devices(project_id)
    
    return render_template('history.html', 
                         history=history, 
                         project=project,
                         devices=devices)

@app.route('/add_maintenance', methods=['POST'])
def add_maintenance():
    project_id = request.form['project_id']
    description = request.form['description']
    manager.add_maintenance_record(project_id, description)
    flash('维护记录已添加！')
    return redirect(url_for('index'))

@app.route('/add_issue', methods=['POST'])
def add_issue():
    project_id = request.form['project_id']
    description = request.form['description']
    manager.add_issue_record(project_id, description)
    flash('故障记录已添加！')
    return redirect(url_for('index'))

@app.route('/update_project_stage/<int:project_id>', methods=['POST'])
def update_project_stage(project_id):
    data = request.get_json()
    new_stage = data.get('stage')
    if new_stage:
        manager.update_project_stage(project_id, new_stage)
        return jsonify({'success': True})
    return jsonify({'success': False, 'error': 'Invalid stage'})

@app.route('/update_project_status/<int:project_id>', methods=['POST'])
def update_project_status(project_id):
    try:
        data = request.get_json()
        new_status = data.get('status')
        add_history = data.get('add_history', False)
        
        if add_history:
            # 获取旧状态用于历史记录
            old_status = manager.get_project_status(project_id)
            success = manager.update_project_status(project_id, new_status, old_status)
        else:
            success = manager.update_project_status(project_id, new_status)
            
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error updating project status: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/reports')
def reports():
    return render_template('reports.html')

@app.route('/get_daily_report/<date>')
@login_required
@permission_required('reports', 'view')
def get_daily_report(date):
    report = manager.get_daily_report(date, session['user_id'])
    return jsonify({
        'success': True,
        'report': report
    })

@app.route('/get_monthly_report/<int:year>/<int:month>')
@login_required
@permission_required('reports', 'view')
def get_monthly_report(year, month):
    reports = manager.get_monthly_report(year, month, session['user_id'])
    return jsonify({
        'success': True,
        'reports': reports  # 直接返回报告列表
    })

@app.route('/delete_project/<int:project_id>', methods=['POST'])
def delete_project(project_id):
    try:
        success = manager.delete_project(project_id)
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error deleting project: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/projects_by_state/<state>')
@login_required
def projects_by_state(state):
    try:
        print(f"\n开始获取{state}状态的项目列表...")
        if state == 'total':
            # 获取所有项目
            projects = manager.get_all_projects()
            print(f"获取到所有项目：{len(projects)} 个")
            for p in projects:
                print(f"项目：ID={p['id']}, 名称={p['client_name']}, 状态={p['state']}")
        else:
            # 获取指定状态的项目
            projects = manager.get_projects_by_state(state)
            print(f"获取到 {state} 状态的项目：{len(projects)} 个")
            for p in projects:
                print(f"项目：ID={p['id']}, 名称={p['client_name']}, 状态={p['state']}")
        
        print("准备渲染模板...")
        return render_template('all_projects.html', 
                             projects=projects)
    except Exception as e:
        print(f"获取项目列表时出错：{str(e)}")
        import traceback
        print(f"错误堆栈：\n{traceback.format_exc()}")
        flash('获取项目列表时出错，请重试')
        return redirect(url_for('index'))

@app.route('/projects_by_stage/<stage>')
def projects_by_stage(stage):
    project_stats = manager.get_project_statistics()
    if stage in project_stats['stages']:
        projects = project_stats['stages'][stage]['projects']
        return render_template('projects_list.html', 
                             projects=projects,
                             state=f'环节：{stage}')
    return redirect(url_for('index'))

@app.route('/add_device_info/<int:project_id>', methods=['POST'])
def add_device_info(project_id):
    try:
        device_info = {
            'type': request.form['device_type'],
            'name': request.form['device_name'],
            'model': request.form['model']
        }
        
        # 根据设备类型处理不同的板卡信息
        if device_info['type'] == '分流设备':
            device_info['cards'] = {
                'mec_10g': int(request.form['mec_10g']),
                'ge_optical': int(request.form['ge_optical']),
                'electrical': int(request.form['electrical'])
            }
        elif device_info['type'] == '数通设备':
            device_info['card_quantity'] = int(request.form['network_card_quantity'])
        elif device_info['type'] == '光旁路保护设':
            device_info['card_quantity'] = int(request.form['obp_card_quantity'])
        
        success = manager.add_device_info(project_id, [device_info])
        if success:
            flash('设备信息添加成功！')
        else:
            flash('设备信息添加失败，请重试')
    except Exception as e:
        flash(f'添加设备信息时出错：{str(e)}')
    
    return redirect(url_for('project_history', project_id=project_id))

@app.route('/add_device/<int:project_id>', methods=['POST'])
def add_device(project_id):
    try:
        data = request.get_json()
        device_info = {
            'type': data['type'],
            'name': data['name'],
            'model': data['model']
        }
        
        if data['type'] == '分流设备':
            device_info['cards'] = {
                'mec_10g': data['cards']['mec_10g'],
                'ge_optical': data['cards']['ge_optical'],
                'electrical': data['cards']['electrical']
            }
        else:
            device_info['card_quantity'] = data['card_quantity']
        
        success = manager.add_device_info(project_id, [device_info])
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error adding device: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/update_device/<int:device_id>', methods=['POST'])
def update_device(device_id):
    try:
        data = request.get_json()
        success = manager.update_device_info(device_id, data)
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error updating device: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/delete_device/<int:device_id>', methods=['POST'])
def delete_device(device_id):
    try:
        success = manager.delete_device(device_id)
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error deleting device: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/update_project_info/<int:project_id>', methods=['POST'])
def update_project_info(project_id):
    try:
        data = request.get_json()
        field = data.get('field')
        value = data.get('value')
        
        # 允许更多字段的更新，包括ID
        allowed_fields = ['id', 'client_name', 'stage', 'status', 'area', 'manager', 'manager_phone']
        if field in allowed_fields:
            success = manager.update_project_info(project_id, field, value)
            return jsonify({'success': success})
        return jsonify({'success': False, 'error': 'Invalid field'})
    except Exception as e:
        print(f"Error updating project info: {e}")
        return jsonify({'success': False, 'error': str(e)}) 

@app.route('/complete_project/<int:project_id>', methods=['POST'])
@login_required
@permission_required('projects', 'edit')
def complete_project(project_id):
    try:
        success = manager.complete_project(project_id)
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error completing project: {e}")
        return jsonify({'success': False, 'error': str(e)}) 

@app.route('/update_daily_report', methods=['POST'])
@login_required
@permission_required('reports', 'edit')
def update_daily_report():
    data = request.get_json()
    success = manager.update_daily_report(
        data['date'], 
        data['content'],
        session['user_id']
    )
    return jsonify({'success': success})

@app.route('/export_daily_report/<date>')
def export_daily_report(date):
    report = manager.get_daily_report(date, session.get('user_id'))
    if report:
        # 创建响应
        response = make_response(report['content'])
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        filename = f'工作日报_{date}.txt'
        encoded_filename = quote(filename)
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        return response
    return '当日无工作记录', 404

@app.route('/export_date_range/<start_date>/<end_date>')
def export_date_range(start_date, end_date):
    user_id = session.get('user_id')
    reports = []
    
    # 将日期范围内的每一天的任务都获取出来
    start = datetime.strptime(start_date, '%Y-%m-%d')
    end = datetime.strptime(end_date, '%Y-%m-%d')
    current = start
    
    while current <= end:
        date_str = current.strftime('%Y-%m-%d')
        daily_report = manager.get_daily_report(date_str, user_id)
        if daily_report and 'content' in daily_report:
            reports.append(daily_report['content'])
        current += timedelta(days=1)
    
    if reports:
        # 合并所有报告内容
        content = '\n\n'.join(reports)
        
        # 创建响应
        response = make_response(content)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        filename = f'工作日报_{start_date}至{end_date}.txt'
        encoded_filename = quote(filename)
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        return response
    
    flash('所选日期范围内没有日报记录')
    return redirect(url_for('index'))

@app.route('/export_monthly_report/<year>/<month>')
def export_monthly_report(year, month):
    user_id = session.get('user_id')
    report = manager.generate_monthly_report(year, month, user_id)
    
    if report:
        # 创建响应
        response = make_response(report)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        filename = f'工作月报_{year}年{month}月.txt'
        encoded_filename = quote(filename)
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        return response
    
    flash('所选月份没有工作记录')
    return redirect(url_for('index'))

@app.route('/export_weekly_report/<year>/<week>')
def export_weekly_report(year, week):
    user_id = session.get('user_id')
    report = manager.generate_weekly_report(year, week, user_id)
    
    if report:
        # 创建响应
        response = make_response(report)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        filename = f'工作周报_{year}年第{week}周.txt'
        encoded_filename = quote(filename)
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        return response
    
    flash('所选周没有工作记录')
    return redirect(url_for('index'))

@app.route('/update_record/<int:record_id>', methods=['POST'])
def update_record(record_id):
    try:
        data = request.get_json()
        description = data.get('description')
        old_value = data.get('old_value')
        new_value = data.get('new_value')
        success = manager.update_record(record_id, description, old_value, new_value)
        return jsonify({
            'success': success,
            'description': description,  # 返���更新后的描述
            'old_value': old_value,     # 返回更新前的值
            'new_value': new_value      # 返回更新后的值
        })
    except Exception as e:
        print(f"Error updating record: {e}")
        return jsonify({'success': False})

@app.route('/delete_record/<int:record_id>', methods=['POST'])
def delete_record(record_id):
    try:
        success = manager.delete_record(record_id)
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error deleting record: {e}")
        return jsonify({'success': False})

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = manager.authenticate_user(username, password)
        if user:
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['role'] = user['role']
            return redirect(url_for('index'))
        flash('用户名或密码错误')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/user_management')
@login_required
@permission_required('users', 'view')
def user_management():
    users = manager.get_users()
    return render_template('user_management.html', users=users)

@app.route('/add_user', methods=['POST'])
@login_required
@permission_required('users', 'add')
def add_user():
    if request.method == 'POST':
        data = request.form
        success = manager.add_user(data)
        if success:
            flash('用户添加成功')
        else:
            flash('用户添加失败')
    return redirect(url_for('user_management'))

@app.route('/toggle_user/<int:user_id>', methods=['POST'])
@login_required
@permission_required('users', 'edit')
def toggle_user(user_id):
    success = manager.toggle_user_status(user_id)
    return jsonify({'success': success})

@app.route('/delete_user/<int:user_id>', methods=['POST'])
@login_required
@permission_required('users', 'delete')
def delete_user(user_id):
    success = manager.delete_user(user_id)
    return jsonify({'success': success})

@app.route('/update_permissions/<int:user_id>', methods=['POST'])
@login_required
@permission_required('users', 'edit')
def update_permissions(user_id):
    data = request.get_json()
    success = manager.update_user_permissions(user_id, data['permissions'])
    return jsonify({'success': success})

@app.route('/get_user_permissions/<int:user_id>')
@login_required
@permission_required('users', 'view')
def get_user_permissions(user_id):
    permissions = manager.get_user_permissions(user_id)
    return jsonify({
        'success': True,
        'permissions': permissions
    })

@app.route('/reactivate_project/<int:project_id>', methods=['POST'])
def reactivate_project(project_id):
    try:
        data = request.get_json()
        state = data.get('state', 'active')  # 默认重新激活为进行中状态
        success = manager.reactivate_project(project_id, state)
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error reactivating project: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/export_project_record/<int:project_id>')
def export_project_record(project_id):
    try:
        content = manager.export_project_record(project_id)
        if content:
            # 创建响应
            response = make_response(content)
            response.headers['Content-Type'] = 'text/plain; charset=utf-8'
            filename = f'项目记录_{project_id}.txt'
            encoded_filename = quote(filename)
            response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
            return response
        else:
            flash('导出失败，项目不存在或无记录')
            return redirect(url_for('project_history', project_id=project_id))
    except Exception as e:
        print(f"Error exporting project record: {e}")
        flash('导出失败，请重试')
        return redirect(url_for('project_history', project_id=project_id))

@app.route('/export_daily_report_word/<date>')
def export_daily_report_word(date):
    user_id = session.get('user_id')
    tasks = manager.get_daily_tasks(date, user_id)
    
    if tasks:
        # 创建Word文档
        doc = Document()
        
        # 设置标题
        title = doc.add_paragraph()
        title_run = title.add_run(f"{date} 工作日报统计")
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # 添加任务内容
        for i, task in enumerate(tasks, 1):
            # 添加任务序号
            doc.add_paragraph(f"{i}. ").add_run()
            
            # 任务类别
            category_para = doc.add_paragraph()
            category_run = category_para.add_run(f"任务类别：{task['project_name']}")
            category_run.font.size = Pt(12)
            
            # 任务内容
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(f"任务内容：{task['content']}")
            content_run.font.size = Pt(12)
            
            # 完成情况
            status_para = doc.add_paragraph()
            status_run = status_para.add_run(f"完成情况：{task['status']}")
            status_run.font.size = Pt(12)
            
            # 添加空行（除了最后一个任务）
            if i < len(tasks):
                doc.add_paragraph()
        
        # 保存到内存中
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # 创建响应
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'工作日报_{date}.docx'
        )
    
    flash('当日无工作记录')
    return redirect(url_for('index'))

@app.route('/edit_project/<int:project_id>', methods=['GET', 'POST'])
def edit_project(project_id):
    if request.method == 'GET':
        try:
            # 获取项目信息
            project = manager.get_project(project_id)
            if not project:
                flash('项目不存在')
                return redirect(url_for('index'))
            
            # 获取区县分组
            district_groups = get_grouped_districts()
            
            return render_template('project_form.html', 
                                 project=project,
                                 district_groups=district_groups,
                                 title='编辑项目')
        except Exception as e:
            print(f"编辑项目时出错：{str(e)}")
            flash('编辑项目时出错，请重试')
            return redirect(url_for('index'))
    else:  # POST 请求
        try:
            # 获取表单数据
            project_data = {
                'client_name': request.form.get('client_name', '').strip(),
                'stage': request.form.get('stage', '').strip(),
                'status': request.form.get('status', '').strip(),
                'area': request.form.get('area', '').strip(),
                'notes': request.form.get('notes', '').strip()
            }
            
            # 验证必填字段
            if not all([project_data['client_name'], project_data['stage'], project_data['area']]):
                missing = []
                if not project_data['client_name']: missing.append('客户名称')
                if not project_data['stage']: missing.append('项目环节')
                if not project_data['area']: missing.append('项目区域')
                flash(f'请填写以下必填字段：{", ".join(missing)}')
                return redirect(url_for('edit_project', project_id=project_id))
            
            # 更新项目
            success = manager.update_project(project_id, project_data)
            if success:
                flash('项目更新成功！')
                return redirect(url_for('project_history', project_id=project_id))
            else:
                flash('项目更新失败，请重试')
                return redirect(url_for('edit_project', project_id=project_id))
                
        except Exception as e:
            print(f"更新项目时出错：{str(e)}")
            flash('更新项目时出错，请重试')
            return redirect(url_for('edit_project', project_id=project_id))